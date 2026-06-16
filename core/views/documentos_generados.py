from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.views.decorators.clickjacking import xframe_options_exempt
from django.contrib import messages
from django.db.models import Q
from django.http import HttpResponse
from ..models import ConfigEmpresa, DocumentoGenerado, Trabajador, Obra, Contrato, Documento
from ..forms import (ConfigEmpresaForm, ContratoTrabajoForm, ContratoTrabajoCreateForm,
                     AnexoContratoForm, FiniquitoForm, PactoHorasExtrasForm,
                     ActaEPPForm, ActaReglamentoForm)

import datetime as _dt
from decimal import Decimal as _Decimal

def _serializable(data):
    """Convert cleaned_data dict to JSON-safe values (date → ISO string, Decimal → float)."""
    result = {}
    for k, v in data.items():
        if isinstance(v, (_dt.date, _dt.datetime)):
            result[k] = v.isoformat()
        elif isinstance(v, _Decimal):
            result[k] = float(v)
        else:
            result[k] = v
    return result


# Mapa: tipo → (FormClass, label)
TIPO_FORMS = {
    'contrato_trabajo': (ContratoTrabajoForm, 'Contrato de Trabajo'),
    'anexo_contrato': (AnexoContratoForm, 'Anexo Contrato de Trabajo'),
    'finiquito': (FiniquitoForm, 'Finiquito Legalizado'),
    'pacto_horas_extras': (PactoHorasExtrasForm, 'Pacto Horas Extraordinarias'),
    'acta_epp': (ActaEPPForm, 'Acta Entrega EPP'),
    'acta_reglamento': (ActaReglamentoForm, 'Acta Entrega Reglamento Interno'),
}


# ── EMPRESAS ────────────────────────────────────────────────────────────────

@login_required
def empresa_list(request):
    empresas = ConfigEmpresa.objects.filter(activo=True)
    form = ConfigEmpresaForm()
    if request.method == 'POST':
        form = ConfigEmpresaForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Empresa creada exitosamente.')
            return redirect('empresa_list')
    return render(request, 'documentos_generados/empresa_list.html', {'empresas': empresas, 'form': form})


@login_required
def empresa_edit(request, pk):
    empresa = get_object_or_404(ConfigEmpresa, pk=pk)
    if request.method == 'POST':
        form = ConfigEmpresaForm(request.POST, instance=empresa)
        if form.is_valid():
            form.save()
            messages.success(request, 'Empresa actualizada.')
            return redirect('empresa_list')
    else:
        form = ConfigEmpresaForm(instance=empresa)
    return render(request, 'documentos_generados/empresa_form.html', {'form': form, 'empresa': empresa})


@login_required
def empresa_delete(request, pk):
    empresa = get_object_or_404(ConfigEmpresa, pk=pk)
    if request.method == 'POST':
        empresa.activo = False
        empresa.save()
        messages.success(request, 'Empresa eliminada.')
    return redirect('empresa_list')


# ── DOCUMENTOS GENERADOS ────────────────────────────────────────────────────

@login_required
def doc_generado_list(request):
    from django.shortcuts import redirect
    return redirect('doc_generado_create')
    # histórico desactivado — lista sin uso mientras no haya documentos generados

@login_required
def _doc_generado_list_historico(request):
    q = request.GET.get('q', '').strip()
    tipo = request.GET.get('tipo', '')
    empresa_id = request.GET.get('empresa_id', '')
    obra_id = request.GET.get('obra_id', '')
    fecha_desde = request.GET.get('fecha_desde', '').strip()
    fecha_hasta = request.GET.get('fecha_hasta', '').strip()

    qs = DocumentoGenerado.objects.filter(activo=True).select_related(
        'trabajador', 'obra', 'empresa'
    ).order_by('-creado_el')

    if q:
        qs = qs.filter(
            Q(trabajador__nombres__icontains=q) |
            Q(trabajador__apellidos__icontains=q) |
            Q(trabajador__rut__icontains=q)
        )
    if tipo:
        qs = qs.filter(tipo=tipo)
    if empresa_id:
        qs = qs.filter(empresa_id=empresa_id)
    if obra_id:
        qs = qs.filter(obra_id=obra_id)
    if fecha_desde:
        try:
            import datetime
            qs = qs.filter(creado_el__date__gte=datetime.date.fromisoformat(fecha_desde))
        except ValueError:
            pass
    if fecha_hasta:
        try:
            import datetime
            qs = qs.filter(creado_el__date__lte=datetime.date.fromisoformat(fecha_hasta))
        except ValueError:
            pass

    empresas = ConfigEmpresa.objects.filter(activo=True)
    obras = Obra.objects.filter(activo=True).order_by('nombre')

    return render(request, 'documentos_generados/list.html', {
        'documentos': qs,
        'q': q,
        'tipo': tipo,
        'empresa_id': empresa_id,
        'obra_id': obra_id,
        'fecha_desde': fecha_desde,
        'fecha_hasta': fecha_hasta,
        'tipo_choices': DocumentoGenerado.TIPO_CHOICES,
        'empresas': empresas,
        'obras': obras,
    })


@login_required
def doc_generado_create(request):
    tipo = request.GET.get('tipo', '')
    trabajador_rut = request.GET.get('trabajador_rut', '')
    obra_id = request.GET.get('obra_id', '')
    contrato_id = request.GET.get('contrato_id', '')
    next_url = request.GET.get('next', '') or request.POST.get('next', '')

    if tipo not in TIPO_FORMS:
        # Show type selection page
        trabajadores = Trabajador.objects.filter(activo=True).order_by('apellidos')
        obras = Obra.objects.filter(activo=True)
        empresas = ConfigEmpresa.objects.filter(activo=True)
        return render(request, 'documentos_generados/seleccionar_tipo.html', {
            'tipo_choices': DocumentoGenerado.TIPO_CHOICES,
            'trabajadores': trabajadores,
            'obras': obras,
            'empresas': empresas,
        })

    FormClass, label = TIPO_FORMS[tipo]
    trabajadores = Trabajador.objects.filter(activo=True)
    obras = Obra.objects.filter(activo=True)
    empresas = ConfigEmpresa.objects.filter(activo=True)

    initial = {}
    trabajador_obj = None
    obra_obj = None
    contrato_obj = None

    if trabajador_rut:
        try:
            trabajador_obj = Trabajador.objects.get(rut=trabajador_rut)
            if tipo == 'contrato_trabajo':
                if trabajador_obj.estado_civil:
                    initial['estado_civil'] = trabajador_obj.estado_civil
                if trabajador_obj.nacionalidad:
                    initial['nacionalidad'] = trabajador_obj.nacionalidad
                if trabajador_obj.procedencia:
                    initial['procedencia'] = trabajador_obj.procedencia
                if trabajador_obj.prevision:
                    initial['prevision'] = trabajador_obj.prevision
                if hasattr(trabajador_obj, 'salud') and trabajador_obj.salud:
                    initial['salud'] = trabajador_obj.salud
                if trabajador_obj.direccion:
                    initial['domicilio_trabajador'] = trabajador_obj.direccion
                if trabajador_obj.ciudad:
                    initial['ciudad_trabajador'] = trabajador_obj.ciudad
        except Trabajador.DoesNotExist:
            pass

    if obra_id:
        try:
            obra_obj = Obra.objects.get(pk=obra_id)
        except Obra.DoesNotExist:
            pass

    # Pre-fill obra en el formulario cuando se viene desde la página masiva
    if obra_obj and tipo == 'contrato_trabajo' and 'obra_contrato' not in initial:
        initial['obra_contrato'] = obra_obj

    if contrato_id:
        try:
            contrato_obj = Contrato.objects.get(pk=contrato_id)
            # Pre-fill labor data from contrato record
            if tipo == 'contrato_trabajo':
                initial['obra_contrato'] = contrato_obj.obra
                initial['especialidad_contrato'] = contrato_obj.especialidad
                initial['tipo_contrato'] = contrato_obj.tipo_contrato
                if contrato_obj.sueldo_base:
                    initial['sueldo_base'] = contrato_obj.sueldo_base
                initial['fecha_inicio_contrato'] = contrato_obj.fecha_inicio
                if contrato_obj.fecha_termino_estimada:
                    initial['fecha_termino_contrato'] = contrato_obj.fecha_termino_estimada
            # Pre-fill personal data from trabajador profile
            if tipo == 'contrato_trabajo' and contrato_obj.trabajador:
                trab = contrato_obj.trabajador
                if trab.estado_civil:
                    initial['estado_civil'] = trab.estado_civil
                if trab.nacionalidad:
                    initial['nacionalidad'] = trab.nacionalidad
                if trab.procedencia:
                    initial['procedencia'] = trab.procedencia
                if trab.prevision:
                    initial['prevision'] = trab.prevision
                if trab.salud:
                    initial['salud'] = trab.salud
                if trab.direccion:
                    initial['domicilio_trabajador'] = trab.direccion
                if trab.ciudad:
                    initial['ciudad_trabajador'] = trab.ciudad
            # For non-contrato_trabajo forms, pre-fill dates/labor from contrato
            if tipo in ('finiquito', 'anexo_contrato'):
                initial['especialidad'] = contrato_obj.especialidad.nombre if contrato_obj.especialidad else ''
                initial['fecha_inicio_contrato'] = contrato_obj.fecha_inicio
                if contrato_obj.fecha_termino_real:
                    initial['fecha_termino_contrato'] = contrato_obj.fecha_termino_real
                elif contrato_obj.fecha_termino_estimada:
                    initial['fecha_termino_contrato'] = contrato_obj.fecha_termino_estimada
            if tipo == 'anexo_contrato':
                initial['fecha_contrato_original'] = contrato_obj.fecha_inicio
                initial['fecha_documento'] = _dt.date.today()
                if contrato_obj.fecha_termino_estimada:
                    initial['nueva_fecha_vigencia'] = contrato_obj.fecha_termino_estimada
            if tipo == 'finiquito':
                if contrato_obj.obra:
                    initial['nombre_obra'] = contrato_obj.obra.nombre
                initial['fecha_documento'] = _dt.date.today()
                _CAUSAL_MAP = {
                    'Mutuo acuerdo': 'Art. 159 N°1 – Mutuo acuerdo',
                    'Despido por necesidades': 'Art. 161 – Necesidades de la empresa',
                    'Término de obra o faena': 'Art. 159 N°5 – Conclusión del trabajo',
                    'Vencimiento plazo': 'Art. 159 N°5 – Conclusión del trabajo',
                }
                if contrato_obj.tipo_termino:
                    initial['causal'] = _CAUSAL_MAP.get(contrato_obj.tipo_termino, '')
        except Contrato.DoesNotExist:
            pass

    # ── Intercepción inteligente: tipo=contrato_trabajo + rut sin contrato_id ──
    # Evita bypasear validaciones (lista negra, re-contratación, finiquito pendiente)
    # Se omite cuando ya viene obra_id (el usuario ya eligió la obra, p.ej. desde imprimir-masivo)
    if tipo == 'contrato_trabajo' and trabajador_obj and not contrato_id and not obra_id and request.method == 'GET':
        if trabajador_obj.en_lista_negra:
            messages.error(
                request,
                f'{trabajador_obj.nombre_completo} está en lista negra y no puede recibir nuevos contratos. '
                f'Motivo: {trabajador_obj.motivo_lista_negra or "sin detalle"}.'
            )
            return redirect('trabajador_detail', rut=trabajador_rut)
        contratos_para_pdf = Contrato.objects.filter(
            trabajador=trabajador_obj,
            estado__in=('Pendiente de Firma', 'Vigente', 'En Licencia', 'Reactivado'),
            activo=True,
        ).select_related('obra', 'especialidad').order_by('-creado_el')
        if contratos_para_pdf.count() == 1:
            from django.urls import reverse
            return redirect(
                reverse('doc_generado_create') +
                f'?tipo=contrato_trabajo&trabajador_rut={trabajador_rut}&contrato_id={contratos_para_pdf.first().pk}'
            )
        obras_disponibles = Obra.objects.filter(
            activo=True, archivada=False
        ).exclude(estado='Cerrada').order_by('nombre')
        return render(request, 'documentos_generados/seleccionar_contrato.html', {
            'trabajador_obj': trabajador_obj,
            'contratos_para_pdf': contratos_para_pdf,
            'obras_disponibles': obras_disponibles,
        })

    # Auto-select empresa: from obra.empresa → then singleton fallback
    empresa_unica = None
    if contrato_obj and contrato_obj.obra and contrato_obj.obra.empresa:
        empresa_unica = contrato_obj.obra.empresa
    elif obra_obj and obra_obj.empresa:
        empresa_unica = obra_obj.empresa
    elif empresas.count() == 1:
        empresa_unica = empresas.first()

    if request.method == 'POST':
        if tipo == 'contrato_trabajo':
            form = ContratoTrabajoCreateForm(request.POST)
        else:
            form = FormClass(request.POST)
        empresa_id = request.POST.get('empresa_id')
        trabajador_rut_post = request.POST.get('trabajador_rut_post', trabajador_rut)

        if form.is_valid() and empresa_id and trabajador_rut_post:
            try:
                empresa = ConfigEmpresa.objects.get(pk=empresa_id)
                trab = Trabajador.objects.get(rut=trabajador_rut_post)
            except (ConfigEmpresa.DoesNotExist, Trabajador.DoesNotExist):
                messages.error(request, 'Empresa o trabajador no válidos.')
                return redirect(request.path + f'?tipo={tipo}')

            if tipo == 'contrato_trabajo':
                # Bloqueo lista negra
                if trab.en_lista_negra:
                    messages.error(
                        request,
                        f'{trab.nombre_completo} está en lista negra y no puede recibir nuevos contratos. '
                        f'Motivo: {trab.motivo_lista_negra or "sin detalle"}.'
                    )
                    return redirect(request.path + f'?tipo={tipo}&trabajador_rut={trab.rut}')
                cd = form.cleaned_data
                # Reusar contrato existente (cualquier estado activo) para evitar duplicados
                contrato_existente = Contrato.objects.filter(
                    trabajador=trab,
                    obra=cd['obra_contrato'],
                    estado__in=('Pendiente de Firma', 'Vigente', 'En Licencia', 'Reactivado'),
                    activo=True,
                ).order_by('-creado_el').first()
                if contrato_existente:
                    contrato_nuevo = contrato_existente
                    contrato_nuevo.especialidad = cd['especialidad_contrato']
                    contrato_nuevo.tipo_contrato = cd['tipo_contrato']
                    contrato_nuevo.sueldo_base = cd['sueldo_base']
                    contrato_nuevo.fecha_inicio = cd['fecha_inicio_contrato']
                    contrato_nuevo.fecha_termino_estimada = cd.get('fecha_termino_contrato')
                    contrato_nuevo.save()
                    messages.info(request, f'Se vinculó al contrato existente #{contrato_existente.pk} (el trabajador ya había sido asignado a esta obra).')
                else:
                    _ESTADOS_TERMINO = ('Finalizado', 'Finiquitado', 'Rescindido', 'Trasladado')
                    contrato_previo = Contrato.objects.filter(
                        trabajador=trab,
                        obra=cd['obra_contrato'],
                        estado__in=_ESTADOS_TERMINO,
                        activo=True,
                    ).order_by('-creado_el').first()
                    contrato_nuevo = Contrato.objects.create(
                        trabajador=trab,
                        obra=cd['obra_contrato'],
                        especialidad=cd['especialidad_contrato'],
                        tipo_contrato=cd['tipo_contrato'],
                        sueldo_base=cd['sueldo_base'],
                        fecha_inicio=cd['fecha_inicio_contrato'],
                        fecha_termino_estimada=cd.get('fecha_termino_contrato'),
                        estado='Pendiente de Firma',
                        es_recontratacion=bool(contrato_previo),
                        contrato_anterior=contrato_previo,
                    )
                    if contrato_previo:
                        tiene_finiquito = Documento.objects.filter(
                            contrato=contrato_previo,
                            tipo_documento__nombre='Finiquito Legalizado',
                            activo=True,
                        ).exists()
                        if not tiene_finiquito:
                            messages.warning(
                                request,
                                f'Re-contratación registrada para {trab.nombre_completo}. '
                                f'El contrato anterior (#{contrato_previo.pk}, {contrato_previo.estado}) '
                                f'no tiene finiquito subido.'
                            )
                        else:
                            messages.info(
                                request,
                                f'{trab.nombre_completo} re-contratado/a. Contrato previo en {cd["obra_contrato"].nombre} ya tiene finiquito.'
                            )
                seccion_a_keys = {'obra_contrato', 'especialidad_contrato', 'tipo_contrato',
                                  'sueldo_base', 'fecha_inicio_contrato', 'fecha_termino_contrato'}
                doc_datos = _serializable({k: v for k, v in cd.items() if k not in seccion_a_keys})
                doc = DocumentoGenerado(
                    tipo=tipo,
                    empresa=empresa,
                    trabajador=trab,
                    contrato=contrato_nuevo,
                    obra=contrato_nuevo.obra,
                    datos=doc_datos,
                    usuario=request.user.username,
                )
            else:
                obra_id_post = request.POST.get('obra_id_post', obra_id)
                contrato_id_post = request.POST.get('contrato_id_post', contrato_id)
                doc = DocumentoGenerado(
                    tipo=tipo,
                    empresa=empresa,
                    trabajador=trab,
                    datos=_serializable(form.cleaned_data),
                    usuario=request.user.username,
                )
                if obra_id_post:
                    try:
                        doc.obra = Obra.objects.get(pk=obra_id_post)
                    except Obra.DoesNotExist:
                        pass
                if contrato_id_post:
                    try:
                        doc.contrato = Contrato.objects.get(pk=contrato_id_post)
                    except Contrato.DoesNotExist:
                        pass
            # Validación finiquito: solo contratos Finalizados
            if tipo == 'finiquito' and doc.contrato:
                if doc.contrato.estado not in ('Finalizado', 'Rescindido'):
                    messages.error(
                        request,
                        f'No se puede generar finiquito para un contrato en estado "{doc.contrato.estado}". '
                        'El contrato debe estar Finalizado o Rescindido primero.'
                    )
                    return redirect(request.path + f'?tipo={tipo}')
            # Desactivar versiones anteriores del mismo tipo para el mismo contrato/trabajador
            _prev = DocumentoGenerado.objects.filter(tipo=tipo, trabajador=trab, activo=True)
            if doc.contrato:
                _prev = _prev.filter(contrato=doc.contrato)
            _prev.update(activo=False)
            doc.save()
            # Anexo de contrato: actualizar fecha_extension en el contrato vinculado
            ext_param = ''
            if tipo == 'anexo_contrato' and doc.contrato:
                nueva_fecha = form.cleaned_data.get('nueva_fecha_vigencia')
                if nueva_fecha:
                    doc.contrato.fecha_extension = nueva_fecha
                    doc.contrato.save(update_fields=['fecha_extension'])
                    ext_param = f'&ext={nueva_fecha.strftime("%d/%m/%Y")}'
                    if next_url:
                        messages.success(
                            request,
                            f'{label} generado. Vigencia extendida hasta {nueva_fecha.strftime("%d/%m/%Y")}. '
                            f'El contrato original no fue modificado.'
                        )
            else:
                if next_url:
                    messages.success(request, f'{label} creado y guardado.')
            if next_url:
                return redirect(next_url)
            return redirect(f'/documentos-empresa/{doc.pk}/preview/?saved=1{ext_param}')
        else:
            if not empresa_id:
                messages.error(request, 'Debe seleccionar una empresa.')
            if not trabajador_rut_post:
                messages.error(request, 'Debe seleccionar un trabajador.')
    else:
        if tipo == 'contrato_trabajo':
            form = ContratoTrabajoCreateForm(initial=initial)
        else:
            form = FormClass(initial=initial)

    return render(request, 'documentos_generados/form.html', {
        'form': form,
        'tipo': tipo,
        'label': label,
        'trabajadores': trabajadores,
        'obras': obras,
        'empresas': empresas,
        'empresa_unica': empresa_unica,
        'trabajador_obj': trabajador_obj,
        'obra_obj': obra_obj,
        'contrato_obj': contrato_obj,
        'trabajador_rut': trabajador_rut,
        'obra_id': obra_id,
        'contrato_id': contrato_id,
        'next_url': next_url,
    })


@login_required
def doc_generado_edit(request, pk):
    doc = get_object_or_404(DocumentoGenerado, pk=pk, activo=True)
    FormClass, label = TIPO_FORMS[doc.tipo]
    # Para contratos vinculados, el label aclara qué se puede editar aquí
    if doc.tipo == 'contrato_trabajo' and doc.contrato:
        label = 'Ajustar datos del PDF — Contrato de Trabajo'
    empresas = ConfigEmpresa.objects.filter(activo=True)
    trabajadores = Trabajador.objects.filter(activo=True)
    obras = Obra.objects.filter(activo=True)

    next_url = request.GET.get('next', '') or request.POST.get('next', '')
    if not next_url and doc.trabajador:
        next_url = f'/trabajadores/{doc.trabajador.rut}/?tab=contratos'

    if request.method == 'POST':
        form = FormClass(request.POST)
        empresa_id = request.POST.get('empresa_id', doc.empresa_id)
        if form.is_valid():
            doc.datos = _serializable(form.cleaned_data)
            if empresa_id:
                try:
                    doc.empresa = ConfigEmpresa.objects.get(pk=empresa_id)
                except ConfigEmpresa.DoesNotExist:
                    pass
            doc.save()
            messages.success(request, f'{label} actualizado.')
            return redirect(next_url) if next_url else redirect('doc_generado_borradores')
    else:
        initial = doc.datos.copy()
        import datetime
        for key, val in initial.items():
            if isinstance(val, str) and len(val) == 10:
                try:
                    initial[key] = datetime.date.fromisoformat(val)
                except ValueError:
                    pass
        form = FormClass(initial=initial)

    return render(request, 'documentos_generados/form.html', {
        'form': form,
        'tipo': doc.tipo,
        'label': label,
        'doc': doc,
        'empresas': empresas,
        'trabajadores': trabajadores,
        'obras': obras,
        'edit_mode': True,
        'next_url': next_url,
    })


def _monto_en_palabras(n):
    n = int(n)
    if n == 0:
        return 'cero'
    UNO = ['', 'un', 'dos', 'tres', 'cuatro', 'cinco', 'seis', 'siete', 'ocho', 'nueve',
           'diez', 'once', 'doce', 'trece', 'catorce', 'quince',
           'dieciséis', 'diecisiete', 'dieciocho', 'diecinueve',
           'veinte', 'veintiún', 'veintidós', 'veintitrés', 'veinticuatro',
           'veinticinco', 'veintiséis', 'veintisiete', 'veintiocho', 'veintinueve']
    DEC = ['', '', '', 'treinta', 'cuarenta', 'cincuenta', 'sesenta', 'setenta', 'ochenta', 'noventa']
    CEN = ['', 'ciento', 'doscientos', 'trescientos', 'cuatrocientos', 'quinientos',
           'seiscientos', 'setecientos', 'ochocientos', 'novecientos']

    def _bajo_mil(x):
        if x < 30:
            return UNO[x]
        elif x < 100:
            d2, u = divmod(x, 10)
            return DEC[d2] + (' y ' + UNO[u] if u else '')
        elif x == 100:
            return 'cien'
        else:
            c, r = divmod(x, 100)
            return CEN[c] + (' ' + _bajo_mil(r) if r else '')

    partes = []
    if n >= 1_000_000:
        m, n = divmod(n, 1_000_000)
        partes.append('un millón' if m == 1 else _bajo_mil(m) + ' millones')
    if n >= 1_000:
        t, n = divmod(n, 1_000)
        partes.append('mil' if t == 1 else _bajo_mil(t) + ' mil')
    if n > 0:
        partes.append(_bajo_mil(n))
    return ' '.join(partes)


@xframe_options_exempt
@login_required
def doc_generado_preview(request, pk):
    doc = get_object_or_404(DocumentoGenerado, pk=pk, activo=True)
    template_map = {
        'contrato_trabajo': 'documentos_generados/print/contrato_trabajo.html',
        'anexo_contrato': 'documentos_generados/print/anexo_contrato.html',
        'finiquito': 'documentos_generados/print/finiquito.html',
        'pacto_horas_extras': 'documentos_generados/print/pacto_horas_extras.html',
        'acta_epp': 'documentos_generados/print/acta_epp.html',
        'acta_reglamento': 'documentos_generados/print/acta_reglamento.html',
        'acta_reactivacion': 'documentos_generados/print/acta_reactivacion.html',
    }
    template = template_map.get(doc.tipo, 'documentos_generados/print/contrato_trabajo.html')

    # Convert ISO date strings back to date objects so |date: filter works in templates
    import datetime as _preview_dt
    d = {}
    for k, v in doc.datos.items():
        if isinstance(v, str) and len(v) == 10:
            try:
                d[k] = _preview_dt.date.fromisoformat(v)
            except ValueError:
                d[k] = v
        else:
            d[k] = v

    # URL de retorno: parámetro explícito, o inferido del trabajador/obra del doc
    back_url = request.GET.get('next', '')
    if not back_url and doc.trabajador:
        back_url = f'/trabajadores/{doc.trabajador.rut}/?tab=contratos'

    ctx = {'doc': doc, 'd': d, 'back_url': back_url}

    if doc.tipo == 'contrato_trabajo' and doc.contrato:
        c = doc.contrato
        sueldo = int(c.sueldo_base)
        ctx['sueldo_formateado'] = f"{sueldo:,}".replace(",", ".")
        ctx['sueldo_palabras'] = _monto_en_palabras(sueldo).capitalize()
        ctx['empresa_print'] = (c.obra.empresa if c.obra and c.obra.empresa else doc.empresa)
    else:
        ctx['empresa_print'] = doc.empresa

    if doc.tipo == 'finiquito':
        try:
            total = (float(d.get('monto_feriado', 0) or 0)
                     + float(d.get('monto_indemnizacion', 0) or 0)
                     + float(d.get('otros_montos_monto', 0) or 0))
            d['total'] = total
        except Exception:
            d['total'] = 0

    return render(request, template, ctx)


@login_required
def doc_generado_blank_preview(request):
    """Render a blank (empty-data) version of a document template for printing."""
    from types import SimpleNamespace
    tipo = request.GET.get('tipo', '')
    template_map = {
        'contrato_trabajo': 'documentos_generados/print/contrato_trabajo.html',
        'anexo_contrato': 'documentos_generados/print/anexo_contrato.html',
        'finiquito': 'documentos_generados/print/finiquito.html',
        'pacto_horas_extras': 'documentos_generados/print/pacto_horas_extras.html',
        'acta_epp': 'documentos_generados/print/acta_epp.html',
        'acta_reglamento': 'documentos_generados/print/acta_reglamento.html',
    }
    tipo_labels = dict(DocumentoGenerado.TIPO_CHOICES)
    if tipo not in template_map:
        return redirect('doc_generado_create')

    empresa_id = request.GET.get('empresa_id', '')
    empresa = None
    if empresa_id:
        try:
            empresa = ConfigEmpresa.objects.get(pk=int(empresa_id), activo=True)
        except (ConfigEmpresa.DoesNotExist, ValueError):
            pass
    if not empresa:
        empresa = ConfigEmpresa.objects.filter(activo=True).first()

    doc_mock = SimpleNamespace(
        pk=None,
        tipo=tipo,
        trabajador=SimpleNamespace(
            nombre_completo='', rut='',
            fecha_nacimiento=None,
            direccion='', telefono='', correo='',
        ),
        empresa=empresa,
        contrato=None,
        contrato_id=None,
        get_tipo_display=lambda: tipo_labels.get(tipo, tipo),
    )

    ctx = {
        'doc': doc_mock,
        'd': {},
        'is_blank': True,
        'tipo_label': tipo_labels.get(tipo, tipo),
        'empresa_print': empresa,
    }
    return render(request, template_map[tipo], ctx)


@login_required
def doc_generado_borradores(request):
    tipo_f = request.GET.get('tipo', '')
    obra_f = request.GET.get('obra_id', '')

    # Limpieza automática de duplicados: conserva solo el más reciente por tipo+contrato
    from django.db.models import Max
    for tipo_val, _ in DocumentoGenerado.TIPO_CHOICES:
        grupos = (
            DocumentoGenerado.objects
            .filter(activo=True, tipo=tipo_val, contrato__isnull=False)
            .values('contrato_id')
            .annotate(max_id=Max('id'))
        )
        for g in grupos:
            DocumentoGenerado.objects.filter(
                tipo=tipo_val, contrato_id=g['contrato_id'], activo=True
            ).exclude(pk=g['max_id']).update(activo=False)

    qs = (
        DocumentoGenerado.objects
        .filter(activo=True)
        .select_related('trabajador', 'obra', 'contrato', 'contrato__obra', 'empresa')
        .order_by('contrato__obra__nombre', 'obra__nombre', 'trabajador__apellidos', '-creado_el')
    )
    if tipo_f:
        qs = qs.filter(tipo=tipo_f)
    if obra_f:
        from django.db.models import Q
        qs = qs.filter(Q(obra_id=obra_f) | Q(contrato__obra_id=obra_f))

    obras = Obra.objects.filter(activo=True).order_by('nombre')

    # Agrupar por obra para el template
    grupos = {}
    for doc in qs:
        obra_nombre = (
            doc.contrato.obra.nombre if doc.contrato and doc.contrato.obra
            else doc.obra.nombre if doc.obra
            else 'Sin obra asignada'
        )
        if obra_nombre not in grupos:
            grupos[obra_nombre] = []
        grupos[obra_nombre].append(doc)

    return render(request, 'documentos_generados/borradores.html', {
        'grupos': grupos,
        'tipo_choices': DocumentoGenerado.TIPO_CHOICES,
        'tipo_f': tipo_f,
        'obra_f': obra_f,
        'obras': obras,
        'total': sum(len(v) for v in grupos.values()),
    })


@login_required
def doc_generado_delete(request, pk):
    doc = get_object_or_404(DocumentoGenerado, pk=pk, activo=True)
    if request.method == 'POST':
        doc.activo = False
        doc.save()
        # Al eliminar un Anexo, revertir la extensión del contrato si no quedan otros Anexos activos
        if doc.tipo == 'anexo_contrato' and doc.contrato:
            otros_anexos = DocumentoGenerado.objects.filter(
                tipo='anexo_contrato', contrato=doc.contrato, activo=True
            ).exists()
            if not otros_anexos:
                doc.contrato.fecha_extension = None
                doc.contrato.save(update_fields=['fecha_extension'])
        messages.success(request, 'Borrador eliminado.')
        next_url = request.POST.get('next', '')
        return redirect(next_url) if next_url else redirect('doc_generado_borradores')
    return redirect('doc_generado_borradores')


@login_required
def doc_generado_imprimir_masivo(request):
    obras = Obra.objects.filter(activo=True).order_by('nombre')
    trabajadores = Trabajador.objects.filter(activo=True).order_by('apellidos', 'nombres')
    obra_id = request.GET.get('obra_id', '') or request.POST.get('obra_id', '')
    obra_obj = None
    borradores = []
    sin_borrador = []
    borrador_trabajador_ids = set()

    if obra_id:
        try:
            obra_obj = Obra.objects.get(pk=obra_id)
            borradores = list(
                DocumentoGenerado.objects
                .filter(activo=True, tipo='contrato_trabajo')
                .filter(Q(obra_id=obra_id) | Q(contrato__obra_id=obra_id))
                .select_related('trabajador', 'contrato', 'contrato__obra', 'empresa')
                .order_by('trabajador__apellidos')
            )
            borrador_trabajador_ids = set(b.trabajador_id for b in borradores if b.trabajador_id)
            # Trabajadores en dotación sin borrador generado
            borrador_contrato_ids = set(b.contrato_id for b in borradores if b.contrato_id)
            from ..models import Contrato as _Contrato
            sin_borrador = (
                _Contrato.objects
                .filter(obra_id=obra_id, estado='Pendiente de Firma', activo=True)
                .exclude(pk__in=borrador_contrato_ids)
                .select_related('trabajador', 'especialidad')
                .order_by('trabajador__apellidos')
            )
        except Obra.DoesNotExist:
            pass

    if request.method == 'POST':
        pks = request.POST.getlist('doc_ids')
        if not pks:
            messages.error(request, 'Seleccione al menos un borrador para imprimir.')
            return redirect(f'/documentos-empresa/imprimir-masivo/?obra_id={obra_id}')

        docs_qs = (
            DocumentoGenerado.objects
            .filter(pk__in=pks, activo=True, tipo='contrato_trabajo')
            .select_related('trabajador', 'contrato', 'contrato__obra', 'empresa')
        )

        import datetime as _preview_dt
        docs_list = []
        for doc in docs_qs:
            d = {}
            for k, v in doc.datos.items():
                if isinstance(v, str) and len(v) == 10:
                    try:
                        d[k] = _preview_dt.date.fromisoformat(v)
                    except ValueError:
                        d[k] = v
                else:
                    d[k] = v

            c = doc.contrato
            sueldo_formateado = ''
            sueldo_palabras = ''
            if c and c.sueldo_base:
                sueldo_int = int(c.sueldo_base)
                sueldo_formateado = f"{sueldo_int:,}".replace(",", ".")
                sueldo_palabras = _monto_en_palabras(sueldo_int).capitalize()

            empresa_print = (c.obra.empresa if c and c.obra and c.obra.empresa else doc.empresa)

            docs_list.append({
                'doc': doc,
                'd': d,
                'sueldo_formateado': sueldo_formateado,
                'sueldo_palabras': sueldo_palabras,
                'empresa_print': empresa_print,
            })

        return render(request, 'documentos_generados/print/masivo_contratos.html', {
            'docs_list': docs_list,
            'obra_obj': obra_obj,
        })

    return render(request, 'documentos_generados/imprimir_masivo.html', {
        'obras': obras,
        'trabajadores': trabajadores,
        'obra_obj': obra_obj,
        'obra_id': obra_id,
        'borradores': borradores,
        'sin_borrador': sin_borrador,
        'borrador_trabajador_ids': borrador_trabajador_ids,
    })


@login_required
def doc_generado_firmar(request, pk):
    """Marca el contrato asociado como Vigente (confirma firma). Solo válido para contrato_trabajo."""
    doc = get_object_or_404(DocumentoGenerado, pk=pk, activo=True)
    if request.method == 'POST' and doc.tipo == 'contrato_trabajo' and doc.contrato and doc.contrato.estado == 'Pendiente de Firma':
        doc.contrato.estado = 'Vigente'
        doc.contrato.save(update_fields=['estado'])
        messages.success(request, f'Contrato #{doc.contrato.pk} marcado como Vigente.')
    else:
        messages.error(request, 'Esta acción solo está disponible para borradores de Contrato de Trabajo.')
    return redirect('doc_generado_list')


@login_required
def doc_generado_pdf_download(request, pk):
    """Genera y descarga el contrato como PDF individual usando xhtml2pdf."""
    from xhtml2pdf import pisa
    from django.template.loader import render_to_string
    from io import BytesIO
    import datetime as _dt

    doc = get_object_or_404(DocumentoGenerado, pk=pk, activo=True)
    if doc.tipo != 'contrato_trabajo':
        messages.error(request, 'La descarga PDF solo está disponible para Contratos de Trabajo.')
        return redirect('doc_generado_preview', pk=pk)

    d = {}
    for k, v in doc.datos.items():
        if isinstance(v, str) and len(v) == 10:
            try:
                d[k] = _dt.date.fromisoformat(v)
            except ValueError:
                d[k] = v
        else:
            d[k] = v

    c = doc.contrato
    sueldo_formateado = ''
    if c and c.sueldo_base:
        sueldo_formateado = f"{int(c.sueldo_base):,}".replace(",", ".")
    empresa_print = (c.obra.empresa if c and c.obra and c.obra.empresa else doc.empresa)

    ctx = {
        'doc': doc,
        'd': d,
        'sueldo_formateado': sueldo_formateado,
        'empresa_print': empresa_print,
    }

    html = render_to_string('documentos_generados/print/pdf_contrato_trabajo.html', ctx, request=request)

    buffer = BytesIO()
    status = pisa.CreatePDF(html.encode('utf-8'), dest=buffer, encoding='utf-8')
    if status.err:
        messages.error(request, 'Error al generar el PDF.')
        return redirect('doc_generado_preview', pk=pk)

    buffer.seek(0)
    apellidos = doc.trabajador.apellidos.replace(' ', '_') if doc.trabajador else 'Trabajador'
    nombres = doc.trabajador.nombres.replace(' ', '_') if doc.trabajador else ''
    nombre_archivo = f"Contrato_{apellidos}_{nombres}.pdf"

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
    return response


@login_required
def doc_generado_word(request, pk):
    """Generate a .docx contract using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    doc_obj = get_object_or_404(DocumentoGenerado, pk=pk, activo=True)
    if doc_obj.tipo != 'contrato_trabajo':
        messages.error(request, 'La generación Word solo está disponible para Contratos de Trabajo.')
        return redirect('doc_generado_preview', pk=pk)

    d = doc_obj.datos
    c = doc_obj.contrato
    trabajador = doc_obj.trabajador

    empresa = (c.obra.empresa if c and c.obra and c.obra.empresa else doc_obj.empresa) if c else doc_obj.empresa

    sueldo_fmt = ''
    sueldo_palabras = ''
    if c and c.sueldo_base:
        sueldo_int = int(c.sueldo_base)
        sueldo_fmt = f"{sueldo_int:,}".replace(",", ".")
        sueldo_palabras = _monto_en_palabras(sueldo_int).capitalize()

    # ── helpers ──
    def _set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _bold_run(para, text, size=None, color=None):
        run = para.add_run(text)
        run.bold = True
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    def _green_run(para, text, size=8):
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.highlight_color = None
        run.font.color.rgb = RGBColor(0, 100, 0)
        return run

    # ── document setup ──
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(2)

    # Default font
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(8.5)

    # ── TÍTULO ──
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CONTRATO DE TRABAJO')
    run.bold = True
    run.font.size = Pt(13)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tipo_txt = (c.tipo_contrato.upper() if c else 'POR OBRA O FAENA')
    p2.add_run(tipo_txt).font.size = Pt(9)

    # Date / city line
    ciudad_doc = d.get('ciudad_documento', '')
    fecha_doc_raw = d.get('fecha_documento', '')
    import datetime
    fecha_doc_str = ''
    if fecha_doc_raw:
        try:
            if isinstance(fecha_doc_raw, str):
                fd = datetime.date.fromisoformat(fecha_doc_raw)
            else:
                fd = fecha_doc_raw
            meses = ['enero','febrero','marzo','abril','mayo','junio',
                     'julio','agosto','septiembre','octubre','noviembre','diciembre']
            fecha_doc_str = f"{fd.day} de {meses[fd.month-1]} de {fd.year}"
        except Exception:
            fecha_doc_str = str(fecha_doc_raw)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.add_run(f"En {ciudad_doc}, a {fecha_doc_str}").font.size = Pt(8.5)

    document.add_paragraph()  # spacer

    # ── TRABAJADOR table ──
    def _hdr_row(table, text):
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[-1]) if len(row.cells) > 1 else None
        _set_cell_bg(cell, 'C6EFCE')
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8.5)

    def _data_row(table, label1, val1, label2='', val2='', green_val=False):
        row = table.add_row()
        cells = row.cells
        cells[0].paragraphs[0].add_run(label1).bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        r = cells[1].paragraphs[0].add_run(str(val1 or ''))
        r.font.size = Pt(8)
        if green_val:
            _set_cell_bg(cells[1], 'C6EFCE')
        if label2:
            cells[2].paragraphs[0].add_run(label2).bold = True
            cells[2].paragraphs[0].runs[0].font.size = Pt(8)
            r2 = cells[3].paragraphs[0].add_run(str(val2 or ''))
            r2.font.size = Pt(8)
            if green_val:
                _set_cell_bg(cells[3], 'C6EFCE')

    # Trabajador section
    tbl = document.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    widths = [Cm(3.8), Cm(5), Cm(3.8), Cm(5)]
    # header
    row0 = tbl.add_row()
    for i, cell in enumerate(row0.cells):
        if i > 0:
            cell.merge(row0.cells[i]) if False else None
    hdr_cell = row0.cells[0]
    for c2 in row0.cells[1:]:
        hdr_cell = hdr_cell.merge(c2)
    _set_cell_bg(hdr_cell, 'C6EFCE')
    hdr_p = hdr_cell.paragraphs[0]
    run = hdr_p.add_run('INDIVIDUALIZACIÓN DEL TRABAJADOR:')
    run.bold = True
    run.font.size = Pt(8.5)

    rows_data = [
        ('Nombre del Trabajador:', trabajador.nombre_completo, 'Nacionalidad:', d.get('nacionalidad', '')),
        ('Cédula Nac. de Identidad:', trabajador.rut, 'Estado Civil:', d.get('estado_civil', '')),
        ('Fecha de Nacimiento:', trabajador.fecha_nacimiento.strftime('%d/%m/%Y') if trabajador.fecha_nacimiento else '', 'Procedencia:', d.get('procedencia', '')),
        ('Domicilio:', d.get('domicilio_trabajador') or trabajador.direccion or '', 'Ciudad:', d.get('ciudad_trabajador', '')),
        ('Previsión:', d.get('prevision', ''), 'Salud:', d.get('salud', '')),
        ('Teléfono:', trabajador.telefono or '', 'Correo:', trabajador.correo or ''),
    ]
    for r_data in rows_data:
        row = tbl.add_row()
        for i, cell in enumerate(row.cells):
            cell.paragraphs[0].add_run(str(r_data[i] or '')).font.size = Pt(8)
            if i % 2 == 0:
                cell.paragraphs[0].runs[0].bold = True

    document.add_paragraph()

    # Empleador section
    tbl2 = document.add_table(rows=0, cols=4)
    tbl2.style = 'Table Grid'
    tbl2.autofit = False

    row0b = tbl2.add_row()
    hdr_cell2 = row0b.cells[0]
    for c2 in row0b.cells[1:]:
        hdr_cell2 = hdr_cell2.merge(c2)
    _set_cell_bg(hdr_cell2, 'C6EFCE')
    run = hdr_cell2.paragraphs[0].add_run('INDIVIDUALIZACIÓN DEL EMPLEADOR:')
    run.bold = True
    run.font.size = Pt(8.5)

    emp_nombre = empresa.nombre if empresa else ''
    emp_rut = empresa.rut_empresa if empresa else ''
    emp_dir = empresa.direccion if empresa else ''
    emp_ciudad = empresa.ciudad if empresa else ''
    emp_rep = empresa.nombre_representante if empresa else ''
    emp_rep_rut = empresa.rut_representante if empresa else ''
    emp_correo = empresa.correo_electronico if empresa and hasattr(empresa, 'correo_electronico') else ''

    emp_rows = [
        ('Nombre del Empleador:', emp_nombre, 'RUT N°', emp_rut),
        ('Dirección:', emp_dir, 'Ciudad:', emp_ciudad),
        ('Representante Legal:', emp_rep, 'RUT N°', emp_rep_rut),
    ]
    if emp_correo:
        emp_rows.append(('Correo Electrónico:', emp_correo, '', ''))

    for r_data in emp_rows:
        row = tbl2.add_row()
        cells = row.cells
        cells[0].paragraphs[0].add_run(r_data[0]).bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        cells[1].paragraphs[0].add_run(str(r_data[1] or '')).font.size = Pt(8)
        _set_cell_bg(cells[1], 'C6EFCE')
        if r_data[2]:
            cells[2].paragraphs[0].add_run(r_data[2]).bold = True
            cells[2].paragraphs[0].runs[0].font.size = Pt(8)
            cells[3].paragraphs[0].add_run(str(r_data[3] or '')).font.size = Pt(8)
            _set_cell_bg(cells[3], 'C6EFCE')
        else:
            # merge last 3 cells for email row
            cells[1].merge(cells[2]).merge(cells[3]) if r_data[2] == '' else None

    document.add_paragraph()

    # Intro paragraph
    p_intro = document.add_paragraph()
    p_intro.add_run(
        'Entre las partes arriba individualizadas se suscribe el presente contrato de trabajo, '
        'para cuyo efecto, los contratantes convienen en denominarse '
    ).font.size = Pt(8.5)
    r_bold = p_intro.add_run('Trabajador y Empleador')
    r_bold.bold = True
    r_bold.font.size = Pt(8.5)
    p_intro.add_run(', respectivamente.').font.size = Pt(8.5)
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── CLÁUSULAS ──
    def _clause(num, text_parts):
        """text_parts: list of (text, bold, green) tuples"""
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        num_run = p.add_run(f'{num}.- ')
        num_run.bold = True
        num_run.font.size = Pt(8.5)
        for text, bold, green in text_parts:
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(8.5)
            if green:
                r.font.color.rgb = RGBColor(0, 100, 0)
        return p

    especialidad = c.especialidad.nombre if c else ''
    obra_nombre = c.obra.nombre if c else ''
    obra_dir = (c.obra.direccion if c and c.obra and c.obra.direccion else '') if c else ''

    _clause('1', [
        ('El Trabajador se compromete a ejecutar la labor de ', False, False),
        (especialidad, True, True),
        (' en la obra denominada ', False, False),
        (obra_nombre, True, True),
        (' ubicada en ', False, False),
        (obra_dir or '_______________', True, True),
        ('; el trabajo se efectuará fundamentalmente en la obra mencionada en el párrafo anterior, '
         'sin perjuicio de la facultad del Empleador para trasladar, por causa justificada, al '
         'trabajador a cualquier Obra que tenga en ejecución dentro de la ciudad, sin que ello '
         'importe menoscabo.', False, False),
    ])

    dias_lab = d.get('dias_laborales', 'Lunes a Viernes')
    horas_sem = d.get('horas_semanales', '44 hrs semanales con 1 hr de colación')
    _clause('2', [
        ('La jornada de trabajo se ejecutará de ', False, False),
        (dias_lab, True, True),
        (' con ', False, False),
        (horas_sem, True, True),
        (' la que no se considerará laborada para ningún efecto legal.', False, False),
    ])

    # Schedule table for clause 2
    inicio_am = d.get('inicio_am', '08:00 HRS')
    termino_am = d.get('termino_am', '13:00 HRS')
    inicio_pm = d.get('inicio_pm', '14:00 HRS')
    termino_lj = d.get('termino_lj', '18:00 HRS')
    termino_v = d.get('termino_v', '17:00 HRS')

    sched = document.add_table(rows=0, cols=5)
    sched.style = 'Table Grid'

    lj_label = 'LUNES A JUEVES'
    v_label = 'VIERNES'
    sched_rows = [
        (lj_label, 'Desde', inicio_am, 'A', termino_am),
        ('', 'Desde', inicio_pm, 'A', termino_lj),
        (v_label, 'Desde', inicio_am, 'A', termino_am),
        ('', 'Desde', inicio_pm, 'A', termino_v),
    ]
    if dias_lab == 'Lunes a Sábado':
        sched_rows += [
            ('SÁBADO', 'Desde', inicio_am, 'A', termino_am),
            ('', '', '', '', 'Tarde libre'),
        ]

    for row_data in sched_rows:
        row = sched.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.paragraphs[0].add_run(str(val)).font.size = Pt(8)
            if i in (0, 2, 4) and val:
                _set_cell_bg(cell, 'C6EFCE')
                if i == 0:
                    cell.paragraphs[0].runs[0].bold = True

    document.add_paragraph()

    tipo_contrato = c.tipo_contrato.upper() if c else '_______________'
    _clause('3', [
        ('El trabajo se ejecutará por: ', False, False),
        (tipo_contrato, True, True),
    ])

    _clause('4', [
        ('El Empleador se compromete a remunerar al trabajador, por mes efectivamente trabajado '
         'en jornada completa, la suma de: $', False, False),
        (sueldo_fmt or '_______________', True, True),
        (' (', False, False),
        (sueldo_palabras or '_______________', True, True),
        (' pesos.-) El tiempo extraordinario se pagará con el recargo legal y se cancelará '
         'conjuntamente con el respectivo sueldo, dentro de los cinco primeros días del mes '
         'siguiente al período vencido, y de cuyo monto el Empleador hará las deducciones que '
         'establecen las leyes vigentes.', False, False),
    ])

    fecha_termino_str = ''
    if c and c.fecha_termino_estimada:
        ft = c.fecha_termino_estimada
        meses = ['enero','febrero','marzo','abril','mayo','junio',
                 'julio','agosto','septiembre','octubre','noviembre','diciembre']
        fecha_termino_str = f"{ft.day} de {meses[ft.month-1]} de {ft.year}"

    _clause('5', [
        ('El presente contrato durará hasta el: ', False, False),
        (fecha_termino_str or '___________________________________', True, True),
        (', sin embargo, podrá ponérsele término cuando concurran para ello causas justificadas, '
         'que en conformidad a la Ley, puedan producir su caducidad o término.', False, False),
    ])

    _clause('6', [
        ('En este acto las partes acuerdan que, sin perjuicio de lo dispuesto en las cláusulas '
         'anteriores, este contrato podrá expirar paulatina y anticipadamente al término final o '
         'entrega definitiva de la obra, en virtud de ir concluyendo en forma decreciente las '
         'labores de la faena. También este contrato terminará anticipadamente cuando el trabajador '
         'incurra en causales para ello, según lo dispuesto en el artículo 160 del Código del '
         'Trabajo, sin derecho a indemnización alguna. Igualmente, este contrato terminará '
         'anticipadamente, si el dueño de la obra termina unilateralmente el contrato con el '
         'empleador, o reduce el monto o cuantía del contrato celebrado, en cuyo caso el trabajador '
         'acepta que la causa de término de la relación laboral, sea el artículo 159, N° 5 del '
         'Código del Trabajo, es decir, "Conclusión del trabajo o servicio que dio origen al '
         'contrato".', False, False),
    ])

    fecha_inicio_str = ''
    if c and c.fecha_inicio:
        fi = c.fecha_inicio
        meses = ['enero','febrero','marzo','abril','mayo','junio',
                 'julio','agosto','septiembre','octubre','noviembre','diciembre']
        fecha_inicio_str = f"{fi.day} de {meses[fi.month-1]} de {fi.year}"

    _clause('7', [
        ('Se deja constancia que el trabajador: ', False, False),
        (trabajador.nombre_completo, True, False),
        (' ingresó al servicio el ', False, False),
        (fecha_inicio_str or '_______________', True, True),
    ])

    _clause('8', [
        ('Se entienden incorporadas al presente contrato todas las disposiciones legales que se '
         'dicten con posterioridad a la fecha de suscripción y que tengan relación con él, y las '
         'cláusulas adicionales estipuladas en el ', False, False),
        ('ANEXO DE CONTRATO DE TRABAJO', True, False),
        ('. -', False, False),
    ])

    _clause('9', [
        ('El presente contrato se firma en tres ejemplares del mismo tenor, dejando expresa '
         'constancia que en este acto el Trabajador recibe uno de ellos.', False, False),
    ])

    observaciones = d.get('observaciones', '')
    if observaciones:
        p_obs = document.add_paragraph()
        p_obs.add_run(observaciones).italic = True

    # ── FIRMAS ──
    document.add_paragraph()
    document.add_paragraph()

    sig_tbl = document.add_table(rows=3, cols=2)
    sig_tbl.style = 'Table Grid'
    # Remove borders — use no border style trick
    for row in sig_tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    sig_tbl.cell(0, 0).paragraphs[0].add_run('')  # spacer
    sig_tbl.cell(0, 1).paragraphs[0].add_run('')

    # Signature lines
    sig_tbl.cell(1, 0).paragraphs[0].add_run('_' * 30)
    sig_tbl.cell(1, 1).paragraphs[0].add_run('_' * 30)

    cell_emp = sig_tbl.cell(2, 0)
    cell_emp.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cell_emp.paragraphs[0].add_run('Firma del Empleador\n')
    r1.bold = True
    r1.font.size = Pt(8)
    cell_emp.paragraphs[0].add_run(f"{emp_nombre}\nRUT: {emp_rut}").font.size = Pt(7)

    cell_trab = sig_tbl.cell(2, 1)
    cell_trab.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cell_trab.paragraphs[0].add_run('Firma del Trabajador\n')
    r2.bold = True
    r2.font.size = Pt(8)
    cell_trab.paragraphs[0].add_run(f"{trabajador.nombre_completo}\nRUT: {trabajador.rut}").font.size = Pt(7)

    # Footer note
    document.add_paragraph()
    p_footer = document.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run(f"Contrato N° {c.pk if c else doc_obj.pk} — Generado por TsDesk")
    r_footer.font.size = Pt(7)
    r_footer.font.color.rgb = RGBColor(150, 150, 150)

    # ── Return as download ──
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    nombre_archivo = f"Contrato_{trabajador.apellidos}_{trabajador.nombres}.docx".replace(' ', '_')
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
    return response
